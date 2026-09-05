import streamlit as st

# 📌 가장 먼저 실행되어야 하는 Streamlit 페이지 설정
st.set_page_config(page_title="CanNest 통합 업무 시스템", layout="wide")

import google.generativeai as genai
import fitz  # PyMuPDF
from PIL import Image, ImageOps
import pillow_heif  # HEIC 지원 라이브러리
import io
import zipfile
import json
import os
import datetime
import time
import uuid
import re 
import urllib.request
from concurrent.futures import ThreadPoolExecutor

# DOCX 생성 라이브러리
try:
    import docx
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    st.error("⚠️ 'python-docx' 라이브러리가 필요합니다. requirements.txt에 python-docx를 추가해 주세요.")

# HEIC 이미지 지원 등록 및 고해상도 제한 해제
pillow_heif.register_heif_opener()
Image.MAX_IMAGE_PIXELS = None

# ==========================================
# 0. Secrets 안전 검사 및 보안 비밀번호 설정
# ==========================================
if "APP_PASSWORD" not in st.secrets or "GEMINI_API_KEY" not in st.secrets:
    st.error("⚠️ Streamlit Cloud의 Secrets 설정이 필요합니다.")
    st.info("우측 하단 [Manage app] -> [Settings] -> [Secrets]에 GEMINI_API_KEY와 APP_PASSWORD를 입력해 주세요.")
    st.stop()

def check_password():
    def password_entered():
        if st.session_state["password"] == st.secrets["APP_PASSWORD"]:
            st.session_state["password_correct"] = True
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        st.title("🔒 CanNest 통합 업무 시스템")
        st.text_input("접속 비밀번호를 입력하세요", type="password", on_change=password_entered, key="password")
        return False
    elif not st.session_state["password_correct"]:
        st.title("🔒 CanNest 통합 업무 시스템")
        st.text_input("접속 비밀번호를 입력하세요", type="password", on_change=password_entered, key="password")
        st.error("비밀번호가 틀렸습니다.")
        return False
    return True

if not check_password():
    st.stop()

# ==========================================
# 1. API 키 및 모델 설정
# ==========================================
GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
genai.configure(api_key=GEMINI_API_KEY)

def safe_generate_content(contents):
    candidate_models = ['gemini-3.6-flash']
    last_error = None
    for model_name in candidate_models:
        try:
            mod = genai.GenerativeModel(model_name)
            response = mod.generate_content(contents)
            return response
        except Exception as e:
            last_error = e
            if "404" in str(e) or "not found" in str(e).lower():
                continue
            else:
                raise e
    raise last_error

# ==========================================
# 2. 실시간 웹페이지 텍스트 크롤러 (URL 분석용)
# ==========================================
def fetch_url_content(url):
    target_url = url.strip()
    if not target_url.startswith("http://") and not target_url.startswith("https://"):
        target_url = "https://" + target_url
    try:
        req = urllib.request.Request(
            target_url, 
            headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'}
        )
        with urllib.request.urlopen(req, timeout=12) as response:
            html = response.read().decode('utf-8', errors='ignore')
            
            # Script 및 Style 태그 제거
            text = re.sub(r'<script\b[^<]*(?:(?!</script>)<[^<]*)*</script>', '', html, flags=re.IGNORECASE)
            text = re.sub(r'<style\b[^<]*(?:(?!</style>)<[^<]*)*</style>', '', html, flags=re.IGNORECASE)
            # HTML 태그 제거
            text = re.sub(r'<[^>]+>', ' ', text)
            # 연속 공백 정제
            text = re.sub(r'\s+', ' ', text).strip()
            return text[:15000]
    except Exception as e:
        return ""

# ==========================================
# 3. 공통 캐싱 및 스마트 글자 크기 조절 함수
# ==========================================
@st.cache_data(ttl=3600, show_spinner=False)
def load_pdf_bytes_cached(file_path):
    if os.path.exists(file_path):
        with open(file_path, "rb") as f:
            return f.read()
    return None

def get_preloaded_file_bytes(file_names):
    for fn in file_names:
        data = load_pdf_bytes_cached(fn)
        if data:
            return data
    return None

def process_uploaded_file_to_image(file_obj):
    if file_obj.type == "application/pdf":
        doc = fitz.open(stream=file_obj.read(), filetype="pdf")
        page = doc.load_page(0)
        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
    else:
        img = Image.open(file_obj)
        img = ImageOps.exif_transpose(img)
        if img.mode != "RGB":
            img = img.convert("RGB")
    
    if img.width > 1500:
        ratio = 1500 / img.width
        new_size = (1500, int(img.height * ratio))
        img = img.resize(new_size, Image.Resampling.LANCZOS)
    
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=70)
    buf.seek(0)
    return Image.open(buf)

def format_full_name(surname, given_name):
    s = surname.strip()
    g = given_name.strip()
    if not s and not g: return ""
    if not s: return g
    if not g: return s
    return f"{g} {s}"

def set_smart_widget_value(widget, value, default_fontsize=11, min_fontsize=5.5):
    val_str = str(value) if value is not None else ""
    widget.field_value = val_str
    
    if hasattr(widget, "field_flags") and widget.field_flags:
        widget.field_flags &= ~1 
        
    try:
        widget.text_font = "Cour"
    except:
        pass
        
    if val_str and hasattr(widget, "rect"):
        box_width = widget.rect.width - 4 
        if box_width > 0:
            try:
                font = fitz.Font("courier") 
                len_at_default = font.text_length(val_str, fontsize=default_fontsize)
                if len_at_default > box_width:
                    len_at_1 = font.text_length(val_str, fontsize=1)
                    if len_at_1 > 0:
                        scaled_size = box_width / len_at_1
                        widget.text_fontsize = max(min_fontsize, min(default_fontsize, scaled_size))
                    else:
                        widget.text_fontsize = default_fontsize
                else:
                    widget.text_fontsize = default_fontsize
            except Exception:
                widget.text_fontsize = default_fontsize
        else:
            widget.text_fontsize = default_fontsize
    else:
        widget.text_fontsize = default_fontsize
        
    widget.update()

def prepare_document_for_gemini(file_bytes, mime_type, file_name=""):
    if "pdf" in mime_type.lower():
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text("text") + "\n"
            
            if len(text.strip()) > 100:
                return [f"\n--- [Document: {file_name}] ---\n{text[:20000]}\n"]
            else:
                images = []
                for page_num in range(min(len(doc), 10)):
                    page = doc.load_page(page_num)
                    pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    buf = io.BytesIO()
                    img.save(buf, format="JPEG", quality=70)
                    buf.seek(0)
                    images.append({"mime_type": "image/jpeg", "data": buf.getvalue()})
                return images
        except:
            pass
    return [{"mime_type": mime_type, "data": file_bytes}]

def batch_process_client_files(client_files):
    def worker(f):
        mime = f.type if f.type else "application/pdf"
        return prepare_document_for_gemini(f.getvalue(), mime, f.name)

    with ThreadPoolExecutor(max_workers=4) as executor:
        results = list(executor.map(worker, client_files))

    flat_contents = []
    for res in results:
        flat_contents.extend(res)
    return flat_contents

def is_minor(dob_str):
    try:
        birth_date = datetime.datetime.strptime(dob_str, "%Y-%m-%d").date()
        today = datetime.date.today()
        age = today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))
        return age < 19
    except:
        return True

# ==========================================
# 4. 오버타임 조항 계산기 (캐나다 퀘벡 제외 전지역)
# ==========================================
def get_provincial_overtime_clause(address_text):
    text = address_text.upper()
    
    if "BC" in text or "BRITISH COLUMBIA" in text:
        return ("Overtime will be paid in accordance with the applicable employment standards legislation, including:\n"
                "1.5 times the employee’s regular wage for hours worked over 8 hours/day or 40 hours/week; and\n"
                "2 times the employee’s regular wage for hours over 12 hours/day")
    elif "AB" in text or "ALBERTA" in text:
        return "1.5 times the employee's regular rate of pay for hours in excess of 8 hours/day or 44 hours/week"
    elif "ON" in text or "ONTARIO" in text or "NB" in text or "NEW BRUNSWICK" in text:
        return "1.5 times the employee's regular rate of pay for hours worked in excess of 44 hours per week"
    elif "SK" in text or "SASKATCHEWAN" in text or "MB" in text or "MANITOBA" in text or "NL" in text or "NEWFOUNDLAND" in text:
        return "1.5 times the employee's regular rate of pay for hours worked over 8 hours/day or 40 hours/week"
    elif "NS" in text or "NOVA SCOTIA" in text or "PE" in text or "PRINCE EDWARD" in text:
        return "1.5 times the employee's regular rate of pay for hours worked in excess of 48 hours per week"
    elif "YT" in text or "YUKON" in text or "NT" in text or "NORTHWEST" in text or "NU" in text or "NUNAVUT" in text:
        return "1.5 times the employee's regular rate of pay for hours worked over 8 hours/day or 40 hours/week"
    else:
        return "Overtime will be paid in accordance with the applicable provincial employment standards legislation for hours worked in excess of standard full-time limits."

# ==========================================
# 5. 잡오퍼 DOCX 생성 엔진
# ==========================================
def generate_job_offer_docx(data):
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # 0. 로고 삽입 (업로드된 로고가 있을 경우)
    logo_bytes = data.get('logo_bytes')
    if logo_bytes:
        try:
            image_stream = io.BytesIO(logo_bytes)
            doc.add_picture(image_stream, width=Inches(2.2))
            doc.add_paragraph()
        except Exception:
            pass

    # 1. 헤더 (고용주 정보)
    p_head = doc.add_paragraph()
    p_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    emp_name_val = data.get('employer_name', '')
    run_comp = p_head.add_run(f"{emp_name_val}\n")
    run_comp.bold = True
    run_comp.font.size = Pt(13)
    
    if data.get('employer_address'):
        p_head.add_run(f"{data.get('employer_address')}\n")
    if data.get('employer_phone'):
        p_head.add_run(f"T. {data.get('employer_phone')}\n")
    if data.get('employer_email'):
        p_head.add_run(f"E. {data.get('employer_email')}\n")
        
    doc.add_paragraph() 
    
    # 2. 날짜 및 수신자
    doc.add_paragraph(data.get('offer_date', datetime.date.today().strftime("%B %d, %Y")))
    doc.add_paragraph(f"Dear {data.get('client_name', 'Employee')},\n")
    
    # 3. 본문 서두
    intro_p = doc.add_paragraph()
    intro_p.add_run("We are pleased to offer you a full-time position as ")
    run_job = intro_p.add_run(f"{data.get('job_title', '')}")
    run_job.bold = True
    intro_p.add_run(f" for a {data.get('employment_term', '3-year')} term with ")
    run_emp = intro_p.add_run(f"{emp_name_val}")
    run_emp.bold = True
    intro_p.add_run(" based on the following terms and conditions:")
    
    # 4. 세부 조건
    doc.add_heading("Terms of Employment", level=2)
    doc.add_paragraph(f"This is a full-time, {data.get('employment_term', '3-year')} employment term starting from the date agreed upon by the employer and employee.")
    
    doc.add_heading("Start Date", level=2)
    doc.add_paragraph(data.get('start_date', 'The employment start date will be as soon as possible upon the employee’s authorization to work in Canada.'))
    
    doc.add_heading("Job Location", level=2)
    doc.add_paragraph(data.get('job_location', ''))
    
    doc.add_heading(f"Job Title: {data.get('job_title', '')}", level=2)
    
    duties_list = data.get('job_duties', [])
    if isinstance(duties_list, str):
        duties_list = [d.strip() for d in duties_list.split('\n') if d.strip()]
        
    if duties_list:
        doc.add_paragraph("Job Duties:")
        for duty in duties_list:
            clean_duty = re.sub(r'^[•\-\*]\s*', '', duty)
            doc.add_paragraph(clean_duty, style='List Bullet')
            
    doc.add_heading("Compensation", level=2)
    doc.add_paragraph(f"The employee will be paid ${data.get('wage', '0.00')} per hour, based on a minimum of {data.get('hours', '30')} hours per week.")
    
    doc.add_heading("Overtime Rate", level=2)
    doc.add_paragraph(data.get('overtime_clause', ''))
    
    doc.add_heading("Benefits", level=2)
    doc.add_paragraph(data.get('benefits', '4% vacation pay'))
    
    doc.add_heading("Confidentiality", level=2)
    doc.add_paragraph(
        f"By accepting the terms of this offer, the employee agrees to keep all confidential information obtained "
        f"during their employment with {emp_name_val} strictly confidential. The employee further agrees that, "
        f"upon termination of employment for any reason, they will return all physical and digital property belonging to or "
        f"originating from {emp_name_val} within five days of receiving notice of termination."
    )
    
    # 5. 종결 및 서명란
    doc.add_paragraph(
        f"We are pleased to extend this offer of employment to you on behalf of {emp_name_val}. "
        f"We are confident that you will make a valuable contribution to our company, and we look forward to working with you."
    )
    
    doc.add_paragraph("\nSincerely,")
    doc.add_paragraph("_________________________________")
    
    p_sig = doc.add_paragraph()
    run_s_name = p_sig.add_run(f"{data.get('signer_name', '')}\n")
    run_s_name.bold = True
    p_sig.add_run(f"{data.get('signer_title', 'Director')}\n")
    p_sig.add_run(f"{emp_name_val}\n")
    if data.get('employer_phone'):
        p_sig.add_run(f"T. {data.get('employer_phone')}\n")
    if data.get('employer_email'):
        p_sig.add_run(f"E. {data.get('employer_email')}")
        
    doc.add_paragraph("\nI accept the terms of this offer:")
    doc.add_paragraph("_________________________________")
    
    p_acc = doc.add_paragraph()
    run_c_name = p_acc.add_run(f"{data.get('client_name', '')}\n")
    run_c_name.bold = True
    if data.get('client_dob'):
        p_acc.add_run(f"Date of Birth: {data.get('client_dob')}")
        
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ==========================================
# 6. 기존 잡오퍼 파싱 함수 (연장 모드용)
# ==========================================
def parse_existing_job_offer(file_bytes, mime_type):
    prompt = """
    Analyze this existing Job Offer document carefully.
    Extract the following details into exact JSON:
    - client_name: Employee name
    - client_dob: Date of birth (YYYY-MM-DD)
    - employer_name: Full Company / Employer Name (Include legal name and dba if present, e.g. 'Agape Sushi Inc. dba Hiro Japan Sushi Xpress')
    - signer_name: Name of Director or Signer
    - signer_title: Title of Signer
    - employer_address: Full address of employer
    - employer_phone: Telephone
    - employer_email: Contact email
    - job_title: Job title
    - wage: Hourly wage (e.g. '20.15')
    - hours: Weekly hours (e.g. '30-40')
    - job_location: Job location address
    - benefits: Benefits text
    - job_duties: Array of bullet point job duties

    Return ONLY a raw valid JSON object.
    """
    contents = prepare_document_for_gemini(file_bytes, mime_type, "Existing_Job_Offer.pdf")
    contents.insert(0, prompt)
    try:
        response = safe_generate_content(contents)
        clean_text = response.text.strip().replace('```json', '').replace('```', '')
        return json.loads(clean_text)
    except Exception as e:
        st.error(f"기존 잡오퍼 파싱 오류: {e}")
        return {}

# ==========================================
# 7. PDF 서식 채우기 & CRM 압축 엔진
# ==========================================
def extract_imm5476_info(image):
    prompt = """
    Analyze this identity document (passport/permit/visa) carefully.
    Extract the following details into exact JSON structure:
    - surname: Family name converted to Title Case (First letter capitalized, e.g., 'KIM' -> 'Kim')
    - given_name: Given names converted to Title Case (e.g., 'EUN SUN' -> 'Eun Sun')
    - dob: Date of birth in YYYY-MM-DD format
    - uci: UCI numbers only if present (10 digits or 8 digits), else empty string
    Return ONLY raw valid JSON object without markdown or code formatting.
    """
    try:
        response = safe_generate_content([prompt, image])
        clean_text = response.text.strip().replace('```json', '').replace('```', '')
        return json.loads(clean_text)
    except Exception as e:
        st.error(f"정보 추출 오류: {e}")
        return None

def extract_all_passports_batch(has_non_acc, images):
    prompt = f"""
    You are an expert OCR system specialized in international passports.
    I am providing {len(images)} passport image(s) in exact order.
    For EACH passport image, extract:
    - surname: Surname / Family name converted to Title Case
    - given_name: Given name(s) converted to Title Case
    - dob: Date of birth in YYYY-MM-DD format
    - passport_number: Passport number in uppercase alphanumeric
    - gender: Sex of the person, strictly "F" or "M"

    Return ONLY a raw valid JSON object:
    {{
      "non_accompanying_parent": {{ "surname": "...", "given_name": "...", "dob": "YYYY-MM-DD", "passport_number": "...", "gender": "M" }} or null,
      "family_members": [ {{ "surname": "...", "given_name": "...", "dob": "YYYY-MM-DD", "passport_number": "...", "gender": "F" }} ]
    }}
    """
    contents = [prompt] + images
    try:
        response = safe_generate_content(contents)
        clean_text = response.text.strip().replace('```json', '').replace('```', '')
        return json.loads(clean_text)
    except Exception as e:
        st.error(f"여권 일괄 추출 오류: {e}")
        return None

def extract_case_prep_info(tmpl_bytes, client_files):
    prompt = """
    You are an expert Canadian immigration case prep assistant.
    The FIRST document is a BLANK reference IRCC IMM form template.
    The REMAINING attached documents are client materials.
    Carefully scan ALL attached client documents to extract all required information matching the IMM form fields.
    Return ONLY a raw valid JSON object with sections and fields.
    """
    contents = [prompt]
    contents.extend(prepare_document_for_gemini(tmpl_bytes, "application/pdf", "Blank_IMM_Form.pdf"))
    contents.extend(batch_process_client_files(client_files))

    try:
        response = safe_generate_content(contents)
        clean_text = response.text.strip().replace('```json', '').replace('```', '')
        return json.loads(clean_text)
    except Exception as e:
        st.error(f"서류 정리 오류: {e}")
        return None

def fill_imm5476(template_bytes, data):
    doc = fitz.open(stream=template_bytes, filetype="pdf")
    target_data = {
        "surname": data.get("surname", ""), "given": data.get("given_name", ""),
        "dob": data.get("dob", ""), "email": data.get("email", ""),
        "address_phone": data.get("address_phone", ""),
        "uci": data.get("uci", "").replace("-", ""), "signDate": data.get("signDate", "")  
    }
    flags = {"surname": False, "given": False, "dob": False, "email": False, "address_phone": False, "uci": False}
    page_date_counters = {}
    
    for page_idx, page in enumerate(doc):
        page_date_counters[page_idx] = 0
        for widget in page.widgets():
            field_name = widget.field_name
            if not field_name: continue
            fname_lower = field_name.lower()
            
            val_to_set = None
            if "family name" in fname_lower and not flags["surname"]:
                val_to_set = target_data["surname"]; flags["surname"] = True
            elif "given name" in fname_lower and not flags["given"]:
                val_to_set = target_data["given"]; flags["given"] = True
            elif "date of birth" in fname_lower and not flags["dob"]:
                val_to_set = target_data["dob"]; flags["dob"] = True
            elif "email" in fname_lower and not flags["email"]:
                val_to_set = target_data["email"]; flags["email"] = True
            elif ("telephone" in fname_lower or "address" in fname_lower) and page_idx == 0 and not flags["address_phone"]:
                val_to_set = target_data["address_phone"]; flags["address_phone"] = True
            elif ("uci" in fname_lower or "unique client identifier" in fname_lower) and not flags["uci"]:
                val_to_set = target_data["uci"]; flags["uci"] = True
            elif "date" in fname_lower and "birth" not in fname_lower:
                page_date_counters[page_idx] += 1
                if page_idx == 2 and page_date_counters[page_idx] == 1:
                    val_to_set = target_data["signDate"]
                elif page_idx == 3 and page_date_counters[page_idx] == 1:
                    val_to_set = target_data["signDate"]
                
            if val_to_set is not None:
                set_smart_widget_value(widget, val_to_set, default_fontsize=9)

    output_pdf = io.BytesIO()
    doc.save(output_pdf); doc.close(); output_pdf.seek(0)
    return output_pdf

def fill_consent_letter(template_bytes, data):
    doc = fitz.open(stream=template_bytes, filetype="pdf")
    children = data.get("children", [])
    num_pages_needed = max(1, (len(children) + 2) // 3)
    for _ in range(num_pages_needed - 1): doc.insert_pdf(doc, from_page=0, to_page=0)
        
    for page_num in range(num_pages_needed):
        page = doc[page_num]
        page_children = children[page_num * 3 : (page_num + 1) * 3]
        child_widgets = [("Information about travelling children", "yyyymmdd"), ("1_2", "2_2"), ("1_3", "2_3")]
        
        for widget in page.widgets():
            fname = widget.field_name.strip() if widget.field_name else ""
            if not fname: continue
            fname_lower = fname.lower()
            
            field_type_str = getattr(widget, "field_type_string", "").lower()
            if "check" in field_type_str or "radio" in field_type_str or "check box" in fname_lower or "checkbox" in fname_lower or "alone" in fname_lower:
                try:
                    widget.field_value = "Off" 
                    widget.update()
                except: pass
                continue
            
            val_to_set = None
            if fname == "1": val_to_set = data.get("non_acc_name", "")
            elif fname == "2": val_to_set = data.get("non_acc_address", "")
            elif fname == "3": val_to_set = data.get("non_acc_phone", "")
            elif fname == "email": val_to_set = data.get("non_acc_email", "")
            elif fname == "This child or these children hashave my or our consent to travel with": val_to_set = data.get("acc_name", "")
            elif fname == "Relationship with Children 1": val_to_set = data.get("acc_relationship", "")
            elif fname == "Relationship with Children 2": val_to_set = data.get("acc_passport", "")
            elif fname == "I give my consent for this child to travel to": val_to_set = "Canada"
            elif fname == "2_4" or fname_lower == "to stay with": val_to_set = data.get("acc_name", "")
            elif fname == "At the following addresses 1": val_to_set = data.get("trip_address", "")
            elif fname == "At the following addresses 2": val_to_set = data.get("trip_phone", "")
            elif fname == "email_2": val_to_set = data.get("trip_email", "")
            elif fname == "yyyymmdd_2": val_to_set = data.get("sign_date", "")
            elif fname == "1_4" or "travel date" in fname_lower or fname_lower == "date" or "from" in fname_lower or "departure" in fname_lower or "start date" in fname_lower: 
                val_to_set = data.get("trip_date", "")
            elif fname_lower == "to" or "return" in fname_lower or "end date" in fname_lower: 
                val_to_set = "" 
            else:
                for idx, (name_key, dob_key) in enumerate(child_widgets):
                    if idx < len(page_children):
                        if fname == name_key: val_to_set = page_children[idx].get("name", "")
                        elif fname == dob_key: val_to_set = page_children[idx].get("dob", "")
            
            if val_to_set is not None:
                set_smart_widget_value(widget, val_to_set, default_fontsize=11)

    output_pdf = io.BytesIO()
    doc.save(output_pdf); doc.close(); output_pdf.seek(0)
    return output_pdf

def process_and_compress_file(file_bytes, mime_type, target_filename):
    is_jpeg = target_filename.lower().endswith(('.jpg', '.jpeg'))
    
    if is_jpeg:
        img = Image.open(io.BytesIO(file_bytes))
        img = ImageOps.exif_transpose(img)
        if img.mode != "RGB":
            img = img.convert("RGB")
            
        if img.width > 2400:
            ratio = 2400 / img.width
            img = img.resize((2400, int(img.height * ratio)), Image.Resampling.LANCZOS)
            
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85, optimize=True)
        buf.seek(0)
        return buf.getvalue(), "image/jpeg"
        
    else:
        if "pdf" in mime_type.lower():
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            total_text_len = 0
            
            for page_idx in range(min(len(doc), 5)):
                page = doc.load_page(page_idx)
                text = page.get_text("text").strip()
                total_text_len += len(text)
                if total_text_len > 50:
                    break
            
            if total_text_len > 50:
                doc.close()
                return file_bytes, "application/pdf"
            
            new_doc = fitz.open()
            target_dpi = 150
            quality = 65
            
            for page in doc:
                zoom = target_dpi / 72.0
                pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                img_buf = io.BytesIO()
                img.save(img_buf, format="JPEG", quality=quality, optimize=True)
                img_buf.seek(0)
                
                pdf_page = new_doc.new_page(width=page.rect.width, height=page.rect.height)
                pdf_page.insert_image(pdf_page.rect, stream=img_buf.getvalue())
                
            output_pdf = io.BytesIO()
            new_doc.save(output_pdf, deflate=True, garbage=4)
            new_doc.close()
            doc.close()
            
            compressed_bytes = output_pdf.getvalue()
            if len(compressed_bytes) >= len(file_bytes):
                return file_bytes, "application/pdf"
                
            return compressed_bytes, "application/pdf"
            
        else:
            target_dpi = 150
            quality = 65
            img = Image.open(io.BytesIO(file_bytes))
            img = ImageOps.exif_transpose(img)
            if img.mode != "RGB":
                img = img.convert("RGB")
                
            if img.width > 1800:
                ratio = 1800 / img.width
                img = img.resize((1800, int(img.height * ratio)), Image.Resampling.LANCZOS)
                
            img_buf = io.BytesIO()
            img.save(img_buf, format="JPEG", quality=quality, optimize=True)
            img_buf.seek(0)
            
            new_doc = fitz.open()
            page_width = img.width * 72 / target_dpi
            page_height = img.height * 72 / target_dpi
            pdf_page = new_doc.new_page(width=page_width, height=page_height)
            
            pdf_page.insert_image(pdf_page.rect, stream=img_buf.getvalue())
            
            output_pdf = io.BytesIO()
            new_doc.save(output_pdf)
            new_doc.close()
            
            compressed_bytes = output_pdf.getvalue()
            if len(compressed_bytes) >= len(file_bytes):
                return file_bytes, "application/pdf"
                
            return compressed_bytes, "application/pdf"

# ==========================================
# 8. Streamlit 네비게이션 및 UI 구성
# ==========================================
MENU_1 = "🍁 IMM5476 자동 작성"
MENU_2 = "✈️ 한부모 동의서 자동 작성"
MENU_3 = "📋 IMM서류 정보 정리"
MENU_4 = "🏷️ CRM 파일명 생성 및 묶기/분할"
MENU_5 = "📄 잡오퍼 DOCX 생성기 (DEV)"

st.sidebar.title("🦅 CanNest Tool")
app_mode = st.sidebar.radio("원하시는 업무 도구를 선택하세요", [MENU_1, MENU_2, MENU_3, MENU_4, MENU_5])

# ------------------------------------------
# 메뉴 1: IMM5476 자동 작성
# ------------------------------------------
if app_mode == MENU_1:
    st.title(MENU_1)
    if "extracted_5476" not in st.session_state: st.session_state.extracted_5476 = None

    template_5476_bytes = get_preloaded_file_bytes(["imm5476_template.pdf", "imm5476_template.pdf.pdf"])
    
    if template_5476_bytes:
        st.success("✅ 사내 표준 'IMM5476' 양식이 자동으로 로드되었습니다.")
    else:
        st.error("⚠️ GitHub에 'imm5476_template.pdf' 파일이 없습니다. 수동으로 업로드해 주세요.")
        template_file = st.file_uploader("IMM5476 템플릿 PDF 선택", type=['pdf'], key="template_5476")
        if template_file: template_5476_bytes = template_file.getvalue()

    st.markdown("---")
    client_file = st.file_uploader("1. 손님 여권 또는 퍼밋", type=['jpg', 'jpeg', 'png', 'pdf'], key="client_5476")

    if client_file and st.button("정보 추출하기", use_container_width=True):
        with st.spinner("서류 분석 중입니다. 잠시만 기다려 주세요..."):
            extracted = extract_imm5476_info(process_uploaded_file_to_image(client_file))
            if extracted: st.session_state.extracted_5476 = extracted; st.success("정보 추출이 완료되었습니다.")

    if st.session_state.extracted_5476:
        data = st.session_state.extracted_5476
        c1, c2 = st.columns(2)
        with c1:
            surname = st.text_input("성 (Surname)", data.get("surname", ""))
            dob = st.text_input("생년월일", data.get("dob", ""))
            email = st.text_input("이메일 주소", "")
        with c2:
            given = st.text_input("이름 (Given Name)", data.get("given_name", ""))
            uci = st.text_input("UCI", data.get("uci", ""))
            sign_date = st.date_input("서명날짜", datetime.date.today())

        address_phone = st.text_input("주소 또는 전화번호 (이메일이 없는 미성년자/신청자용)", placeholder="예: 2301-6658 Dow Ave, Burnaby BC V5H 0C7")

        if st.button("문서 생성 및 다운로드", type="primary"):
            if not template_5476_bytes:
                st.error("템플릿 파일이 없습니다.")
            else:
                final_data = {
                    "surname": surname, "given_name": given, "dob": dob, "uci": uci, 
                    "email": email, "address_phone": address_phone, 
                    "signDate": sign_date.strftime("%Y-%m-%d")
                }
                pdf_out = fill_imm5476(template_5476_bytes, final_data)
                st.download_button("📥 다운로드", pdf_out, file_name=f"IMM5476_{surname}_{given}.pdf", mime="application/pdf")
        
        st.markdown("---")
        if st.button("🔄 전체 리셋", type="secondary", use_container_width=True, key="reset_btn_m1"):
            for key in list(st.session_state.keys()):
                if key != "password_correct":
                    del st.session_state[key]
            st.rerun()

# ------------------------------------------
# 메뉴 2: 한부모 동의서 자동 작성
# ------------------------------------------
elif app_mode == MENU_2:
    st.title(MENU_2)
    if "consent_non_acc" not in st.session_state: st.session_state.consent_non_acc = {}
    if "consent_family" not in st.session_state: st.session_state.consent_family = []

    consent_template_bytes = get_preloaded_file_bytes(["consent_template.pdf", "consent_template.pdf.pdf"])

    if consent_template_bytes:
        st.success("✅ 사내 표준 '한부모 동의서' 양식이 자동으로 로드되었습니다.")
    else:
        st.error("⚠️ GitHub에 'consent_template.pdf' 파일이 없습니다. 수동으로 업로드해 주세요.")
        consent_template = st.file_uploader("동의서 양식 선택", type=['pdf'])
        if consent_template: consent_template_bytes = consent_template.getvalue()

    st.markdown("---")
    c1, c2 = st.columns(2)
    with c1: non_acc_file = st.file_uploader("비동반 부모님 여권 (1장)", type=['jpg', 'jpeg', 'png', 'pdf'])
    with c2: family_files = st.file_uploader("동반 부모/자녀 여권", type=['jpg', 'jpeg', 'png', 'pdf'], accept_multiple_files=True)

    if st.button("여권 정보 추출하기", type="primary", use_container_width=True):
        images = []
        has_non_acc = bool(non_acc_file)
        if non_acc_file: images.append(process_uploaded_file_to_image(non_acc_file))
        if family_files: images.extend([process_uploaded_file_to_image(f) for f in family_files])
        
        if images:
            with st.spinner("여권 정보를 분석 중입니다. 잠시만 기다려 주세요..."):
                res = extract_all_passports_batch(has_non_acc, images)
                if res:
                    st.session_state.consent_non_acc = res.get("non_accompanying_parent", {}) or {}
                    st.session_state.consent_family = res.get("family_members", []) or []
                    st.success("여권 정보 추출이 완료되었습니다.")

    non_acc_data = st.session_state.consent_non_acc
    non_acc_name = st.text_input("비동반 부모 성명", format_full_name(non_acc_data.get('surname',''), non_acc_data.get('given_name','')))
    non_acc_address = st.text_input("주소")
    ca, cb = st.columns(2)
    with ca: non_acc_phone = st.text_input("전화번호")
    with cb: non_acc_email = st.text_input("이메일")

    acc_parents = [p for p in st.session_state.consent_family if not is_minor(p.get("dob", ""))]
    children_list = [p for p in st.session_state.consent_family if is_minor(p.get("dob", ""))]
    
    acc_name, acc_passport, acc_rel = "", "", "Mother"
    if acc_parents:
        p = acc_parents[0]
        acc_name = format_full_name(p.get('surname',''), p.get('given_name',''))
        acc_passport = p.get("passport_number", "")
        acc_rel = "Mother" if p.get("gender") == "F" else "Father"

    cp1, cp2, cp3 = st.columns(3)
    with cp1: acc_name = st.text_input("동반 부모 성명", acc_name)
    with cp2: acc_rel = st.selectbox("관계", ["Mother", "Father"], index=0 if acc_rel=="Mother" else 1)
    with cp3: acc_passport = st.text_input("여권번호", acc_passport)

    final_children = []
    for idx, c in enumerate(children_list):
        cc1, cc2 = st.columns(2)
        with cc1: n = st.text_input(f"자녀{idx+1} 성명", format_full_name(c.get('surname',''), c.get('given_name','')))
        with cc2: d = st.text_input(f"자녀{idx+1} 생일", c.get("dob", "").replace("-", "/"))
        final_children.append({"name": n, "dob": d})
    
    if not children_list:
        cc1, cc2 = st.columns(2)
        with cc1: n = st.text_input("자녀 성명")
        with cc2: d = st.text_input("자녀 생일")
        if n: final_children.append({"name": n, "dob": d})

    trip_address = st.text_input("현지 주소")
    ct1, ct2 = st.columns(2)
    with ct1: trip_phone = st.text_input("현지 전화")
    with ct2: trip_email = st.text_input("현지 이메일")
    
    trip_date = st.text_input("여행 기간 (Travel Date)", placeholder="예: 2026/09/01 ~ 2026/09/30 또는 August 2026")
    
    sign_date_str = st.date_input("서명일", datetime.date.today()).strftime("%Y/%m/%d")

    if st.button("문서 생성 및 다운로드", type="primary"):
        if not consent_template_bytes:
            st.error("양식 파일이 없습니다.")
        else:
            data_consent = {
                "non_acc_name": non_acc_name, "non_acc_address": non_acc_address, "non_acc_phone": non_acc_phone, "non_acc_email": non_acc_email,
                "children": final_children, "acc_name": acc_name, "acc_relationship": acc_rel, "acc_passport": acc_passport,
                "trip_address": trip_address, "trip_phone": trip_phone, "trip_email": trip_email, 
                "trip_date": trip_date, 
                "sign_date": sign_date_str
            }
            pdf_out = fill_consent_letter(consent_template_bytes, data_consent)
            
            crm_name = "NAME"
            if non_acc_name:
                if re.search(r'[가-힣]', non_acc_name):
                    crm_name = non_acc_name.replace(" ", "")
                else:
                    parts = non_acc_name.strip().split()
                    if parts: crm_name = parts[0].capitalize()
            
            download_file_name = f"{crm_name}_Consent Letter for Children Travelling Abroad.pdf"
            
            st.download_button("📥 다운로드", pdf_out, download_file_name, "application/pdf")
            
    st.markdown("---")
    if st.button("🔄 전체 리셋", type="secondary", use_container_width=True, key="reset_btn_m2"):
        for key in list(st.session_state.keys()):
            if key != "password_correct":
                del st.session_state[key]
        st.rerun()

# ------------------------------------------
# 메뉴 3: 이민서류 정보 정리 (Case File Prep)
# ------------------------------------------
elif app_mode == MENU_3:
    st.title(MENU_3)
    
    if "prep_result" not in st.session_state:
        st.session_state.prep_result = None

    st.subheader("1. 대상 서식 (IMM PDF)")
    
    form_map = {
        "직접 파일 업로드 (기타 서식)": None,
        "IMM1294 (SP-OUTSIDE)": "imm1294.pdf",
        "IMM1295 (WP-OUTSIDE)": "imm1295.pdf",
        "IMM5708 (VR-INSIDE)": "imm5708.pdf",
        "IMM5709 (SP-INSIDE)": "imm5709.pdf",
        "IMM5710 (WP-INSIDE)": "imm5710.pdf"
    }
    
    selected_form = st.selectbox("📌 템플릿 서식 선택", list(form_map.keys()))
    
    tmpl_bytes = None
    
    if form_map[selected_form] is not None:
        file_path = form_map[selected_form]
        tmpl_bytes = load_pdf_bytes_cached(file_path)
        if tmpl_bytes:
            st.success(f"✅ '{selected_form}' 양식이 자동으로 로드되었습니다.")
        else:
            st.error(f"⚠️ {file_path} 파일이 서버에 없습니다. 파일 업로드 상태를 확인해 주세요.")
    else:
        tmpl_prep_file = st.file_uploader("빈 IMM 서식 (반드시 Print to PDF로 평탄화된 파일)", type=['pdf'], key="case_tmpl")
        if tmpl_prep_file:
            tmpl_bytes = tmpl_prep_file.getvalue()

    st.markdown("---")
    st.subheader("2. 손님 제출 서류 (복수 선택 가능)")
    client_prep_files = st.file_uploader("질문지, 여권, 퍼밋 등 서류 선택", type=['jpg', 'jpeg', 'png', 'pdf'], accept_multiple_files=True, key="case_client_docs")

    if st.button("서류 정보 정리하기", type="primary", use_container_width=True):
        if tmpl_bytes is None:
            st.warning("1번 단계에서 서식이 정상적으로 선택되거나 업로드되지 않았습니다.")
        elif not client_prep_files:
            st.warning("2번 단계에서 손님 서류를 1개 이상 올려주세요.")
        else:
            with st.spinner("서류를 대조하여 정보 및 불일치 항목을 확인 중입니다. 잠시만 기다려 주세요..."):
                res = extract_case_prep_info(tmpl_bytes, client_prep_files)
                if res:
                    st.session_state.prep_result = res
                    st.success("서류 정보 정리가 완료되었습니다.")

    if st.session_state.prep_result:
        st.markdown("---")
        st.subheader("3. 정리된 정보 결과")

        parsed = st.session_state.prep_result
        sections = parsed.get("sections", [])

        if not sections:
            st.error("서식에서 분석할 항목을 찾지 못했습니다. (파일이 평탄화된 PDF인지 확인해 주세요)")
        else:
            full_text_list = []
            for sec in sections:
                sec_name = sec.get("section", "기타 항목")
                st.write(f"### 📌 {sec_name}")
                full_text_list.append(f"[{sec_name}]")

                table_data = []
                prev_group = None
                for f in sec.get("fields", []):
                    field_lbl = f.get("field", "")
                    val = f.get("value", "")
                    src = f.get("source", "")

                    group_match = re.match(r'^(.*?\bEntry\s*\d+)', field_lbl, re.IGNORECASE)
                    curr_group = group_match.group(1).strip() if group_match else None

                    if prev_group and curr_group and prev_group != curr_group:
                        table_data.append({
                            "항목 (Field)": "──────────",
                            "추출값 (Value)": "──────────",
                            "출처 (Source)": "──────────"
                        })
                        full_text_list.append("")

                    prev_group = curr_group

                    if not val:
                        display_val = "⚠️ 확인 필요 (미발견)"
                        full_text_list.append(f"{field_lbl}: (확인 필요)")
                    else:
                        display_val = val
                        full_text_list.append(f"{field_lbl}: {val}")

                    table_data.append({
                        "항목 (Field)": field_lbl,
                        "추출값 (Value)": display_val,
                        "출처 (Source)": src if src else "-"
                    })

                st.table(table_data)
                full_text_list.append("")

            st.markdown("#### 📋 한눈에 복사하기")
            st.text_area("아래 텍스트를 복사하여 서식에 옮겨 적으세요", value="\n".join(full_text_list), height=250)
            
            if st.button("＋ 새 케이스 정리하기"):
                st.session_state.prep_result = None
                st.rerun()

# ------------------------------------------
# 메뉴 4: CRM 파일명 자동 생성 및 묶기/분할
# ------------------------------------------
elif app_mode == MENU_4:
    st.title(MENU_4)
    st.caption("개별 낱장 이미지, 여러 장짜리 통짜 PDF 등을 섞어서 올려도 AI가 알아서 문서 단위로 묶거나 분할하여 CRM 파일명으로 최적화합니다.")

    if "uploader_key" not in st.session_state:
        st.session_state.uploader_key = str(uuid.uuid4())
    if "analysis_results" not in st.session_state:
        st.session_state.analysis_results = None

    uploaded_files = st.file_uploader(
        "서류 업로드 (복수 선택 가능)", 
        type=['jpg', 'jpeg', 'png', 'pdf', 'heic'], 
        accept_multiple_files=True,
        key=st.session_state.uploader_key
    )

    if uploaded_files:
        if st.button("서류 분석 및 묶기/분할 시작", type="primary", use_container_width=True):
            results = []
            status_text = st.empty()
            progress_bar = st.progress(0)
            
            status_text.text("1. 전체 서류 페이지 스캔 및 미리보기 생성 중...")
            
            global_pages = []
            page_counter = 1
            
            for file in uploaded_files:
                file_bytes = file.getvalue()
                mime_type = file.type if file.type else "application/pdf"
                
                if "pdf" in mime_type.lower():
                    doc = fitz.open(stream=file_bytes, filetype="pdf")
                    for i in range(len(doc)):
                        if page_counter > 40: break
                        page = doc.load_page(i)
                        pix = page.get_pixmap(matrix=fitz.Matrix(1.2, 1.2))
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        buf = io.BytesIO()
                        img.save(buf, format="JPEG", quality=60)
                        
                        global_pages.append({
                            "global_idx": page_counter,
                            "original_name": file.name,
                            "mime_type": mime_type,
                            "file_bytes": file_bytes,
                            "pdf_page_idx": i,
                            "preview_bytes": buf.getvalue()
                        })
                        page_counter += 1
                    doc.close()
                else:
                    if page_counter > 40: continue
                    img = Image.open(io.BytesIO(file_bytes))
                    img = ImageOps.exif_transpose(img)
                    if img.mode != "RGB": img = img.convert("RGB")
                    
                    preview = img.copy()
                    if preview.width > 1200:
                        preview = preview.resize((1200, int(preview.height * (1200/preview.width))), Image.Resampling.LANCZOS)
                    buf = io.BytesIO()
                    preview.save(buf, format="JPEG", quality=60)
                    
                    global_pages.append({
                        "global_idx": page_counter,
                        "original_name": file.name,
                        "mime_type": mime_type,
                        "file_bytes": file_bytes,
                        "pdf_page_idx": 0,
                        "preview_bytes": buf.getvalue()
                    })
                    page_counter += 1

            if page_counter > 40:
                st.warning("⚠️ 업로드된 총 페이지 수가 40장을 초과하여, 앞의 40장까지만 분석 및 병합합니다.")

            status_text.text("2. AI가 페이지별 문맥을 분석하여 연관 서류를 묶거나 나누는 중입니다...")
            
            prompt = f"""
            You are an expert AI document classifier for a Canadian immigration firm.
            I am providing {len(global_pages)} pages of documents uploaded by a client. 

            Your task:
            1. Read ALL pages carefully.
            2. GROUP the pages that logically belong to the SAME document type for the SAME client. 
               *CRITICAL MERGE RULE*: If you see multiple pages of BANK STATEMENTS, PAYSTUBS, or UTILITY BILLS for the SAME client (even if from different months), MERGE THEM ALL into a single group.
            3. ROTATION CHECK: Check if text is upside down or sideways (0, 90, 180, or 270).
            4. For EACH grouped document, generate an EXACT filename using our strict CRM manual rules provided below.

            [STRICT CRM MANUAL FILENAME RULES]
            Rule 1 (Name): Korean client -> Full Korean Name without spaces (e.g. 홍길동). Foreign client -> VERY FIRST WORD of English Given Name in Title Case (e.g. Richard, Pham).
            Rule 2 (Dates): MUST be YYYY.MM.DD (e.g. 2022.01.02). Use dots '.' as delimiters.
            Rule 3 (Delimiter): Underscore '_' between Name, Document Category, and Date/Company.
            Rule 4 (Date Range): Use hyphen '-' for ranges (e.g. 2022.04.22-2022.05.22).
            Rule 5 (Capitalization): Every English word MUST be Title Case (Capitalize First Letter).

            [CATEGORIES & CMS SUFFIX RULES]
            1. Passport: {{Name}}_PP_{{ExpiryDate YYYY.MM.DD}}
            2. Work Permit / Study Permit / Visitor Record / Coop / PGWP / BOWP: {{Name}}_{{WP/SP/VR/Coop/PGWP/BOWP}}_{{ExpiryDate YYYY.MM.DD}}
            3. Questionnaire: {{Name}}_QA_{{Type}}_{{ReceivedDate YYYY.MM.DD}}
            4. Police Certificate: {{Name}}_Police Cert_{{CountryInEnglish}}
            5. Employment Letter / Certificate of Employment / Confirmation of Employment: {{Name}}_LOE_{{CompanyInEnglish}} (CRITICAL: MUST use 'LOE', NEVER 'Employment Letter')
            6. Paystub: {{Name}}_Paystub_{{CompanyInEnglish}}_{{StartDate-EndDate}}
            7. Education Certificate / WES: {{Name}}_{{Diploma/Bachelor/Master/Highschool/Certificate/WES}}_{{SchoolName}} (If WES: {{Name}}_WES)
            8. Certificate of Income: {{Name}}_COI_{{Year YYYY}}
            9. Official English Score: {{Name}}_{{IELTS/CELPIP}}_{{Date YYYY.MM.DD}}
            10. Resume: {{Name}}_Resume_{{ReceivedDate YYYY.MM.DD}}
            11. Medical Exam / Emedical: {{Name}}_Emedical_{{Year YYYY}}
            12. Marriage Certificate: {{Name}}_Marriage Cert_{{IssueDate YYYY.MM.DD}}
            13. Transcript: {{Name}}_Transcript_{{SchoolName}}
            14. Bank Statement: {{Name}}_Bank Statement_{{Year YYYY}}
            15. Family Certificate (가족관계증명서): {{Name}}_Family Cert_{{IssueDate YYYY.MM.DD}} (CRITICAL: MUST use 'Family Cert')
            16. Basic Certificate (기본증명서): {{Name}}_Basic Cert
            17. Birth Certificate (출생증명서): {{Name}}_Birth Cert
            18. Travel Consent (한부모동의서): {{ChildName}}_Travel Consent
            19. ECE License: {{Name}}_ECE License_{{Province}}
            20. Letter of Acceptance (입학허가서): {{Name}}_LOA_{{SchoolName}}
            21. Tuition Receipt: {{Name}}_Tuition Receipt_{{SchoolName}}
            22. Confirmation of Enrollment: {{Name}}_Confirmation of Enrollment_{{SchoolName}}
            23. Digital Photo / Passport Photo: {{Name}}_Digital Photo.jpg

            [CRITICAL FALLBACK RULE FOR UNKNOWN DOCUMENTS]
            - Step 1: If a document does NOT match any of the 23 categories above, extract the official document title printed at the top of the document (in English, Title Case) and format as: {{Name}}_{{DocumentTitleInEnglish}}.
            - Step 2: If the document title/type is completely ambiguous, unreadable, or unclassified, set suggested_filename as {{Name}}_Unclassified_확인필요.pdf and set "is_unclassified": true.

            Return ONLY a raw JSON object:
            {{
              "rotations": {{
                "1": 0,
                "2": 90
              }},
              "documents": [
                {{
                  "client_name": "...",
                  "doc_category": "...",
                  "suggested_filename": "...",
                  "page_indices": [1, 2, 3],
                  "is_unclassified": false
                }}
              ]
            }}
            """
            
            contents = [prompt]
            for p in global_pages:
                contents.append(f"--- Page {p['global_idx']} ---")
                contents.append({"mime_type": "image/jpeg", "data": p['preview_bytes']})
            
            try:
                response = safe_generate_content(contents)
                clean_text = response.text.strip().replace('```json', '').replace('```', '')
                data = json.loads(clean_text)
                raw_docs_info = data.get("documents", [])
                rotations = data.get("rotations", {})
                
                docs_info = []
                for d in raw_docs_info:
                    c_name = d.get("client_name", "").strip()
                    c_cat = d.get("doc_category", "").strip()
                    
                    existing = None
                    for item in docs_info:
                        name1 = item.get("suggested_filename", "").replace(".pdf", "").replace(".jpg", "").strip()
                        name2 = d.get("suggested_filename", "").replace(".pdf", "").replace(".jpg", "").strip()
                        if name1 == name2:
                            existing = item
                            break
                        if c_name and c_cat and c_name == item.get("client_name", "").strip() and c_cat == item.get("doc_category", "").strip():
                            if "bank statement" in c_cat.lower() or "paystub" in c_cat.lower() or "statement" in c_cat.lower():
                                existing = item
                                break
                                
                    if existing:
                        existing["page_indices"] = sorted(list(set(existing.get("page_indices", []) + d.get("page_indices", []))))
                    else:
                        docs_info.append(d)
                        
            except Exception as e:
                st.error(f"AI 분석 중 오류가 발생했습니다: {e}")
                docs_info = []
                rotations = {}

            status_text.text("3. 분석된 정보를 바탕으로 최종 서류 결합 및 회전/압축 중입니다...")
            progress_step = 1 / max(len(docs_info), 1)

            failed_docs = []

            for idx, doc_info in enumerate(docs_info):
                indices = doc_info.get("page_indices", [])
                if not indices: continue
                
                final_name = doc_info.get("suggested_filename", f"Document_{idx+1}")
                is_unclassified = doc_info.get("is_unclassified", False) or "확인필요" in final_name
                
                if not (final_name.lower().endswith(".pdf") or final_name.lower().endswith(".jpg") or final_name.lower().endswith(".jpeg")):
                    if "photo" in final_name.lower() and len(indices) == 1:
                        final_name += ".jpg"
                    else:
                        final_name += ".pdf"
                        
                source_names = []
                group_pages = []
                for p_idx in indices:
                    p_data = next((p for p in global_pages if p['global_idx'] == p_idx), None)
                    if p_data:
                        group_pages.append(p_data)
                        if p_data["original_name"] not in source_names:
                            source_names.append(p_data["original_name"])

                unique_src_files = list(set([p["original_name"] for p in group_pages]))
                is_all_from_same_pdf = (len(unique_src_files) == 1 and "pdf" in group_pages[0]["mime_type"].lower())
                
                needs_rotation = any(int(rotations.get(str(p["global_idx"]), 0)) != 0 for p in group_pages)
                
                try:
                    if is_all_from_same_pdf:
                        src_doc = fitz.open(stream=group_pages[0]["file_bytes"], filetype="pdf")
                        pdf_indices = [p["pdf_page_idx"] for p in group_pages]
                        
                        if len(pdf_indices) == len(src_doc) and not needs_rotation and pdf_indices == list(range(len(src_doc))):
                            merged_pdf_bytes = group_pages[0]["file_bytes"]
                            src_doc.close()
                        else:
                            src_doc.select(pdf_indices)  
                            for i, p_data in enumerate(group_pages):
                                try:
                                    rot = int(rotations.get(str(p_data["global_idx"]), 0))
                                    if rot != 0:
                                        page = src_doc[i]
                                        page.set_rotation((page.rotation + rot) % 360)
                                except: pass
                            merged_pdf_bytes_io = io.BytesIO()
                            src_doc.save(merged_pdf_bytes_io, deflate=True) 
                            src_doc.close()
                            merged_pdf_bytes = merged_pdf_bytes_io.getvalue()
                    else:
                        new_doc = fitz.open()
                        for p_data in group_pages:
                            rot = 0
                            try: rot = int(rotations.get(str(p_data["global_idx"]), 0))
                            except: pass
                            
                            if "pdf" in p_data["mime_type"].lower():
                                src_doc = fitz.open(stream=p_data["file_bytes"], filetype="pdf")
                                new_doc.insert_pdf(src_doc, from_page=p_data["pdf_page_idx"], to_page=p_data["pdf_page_idx"])
                                if rot != 0:
                                    page = new_doc[-1]
                                    page.set_rotation((page.rotation + rot) % 360)
                                src_doc.close()
                            else:
                                img = Image.open(io.BytesIO(p_data["file_bytes"]))
                                img = ImageOps.exif_transpose(img) 
                                if img.mode != "RGB": img = img.convert("RGB")
                                if rot != 0: img = img.rotate(-rot, expand=True) 
                                img_buf = io.BytesIO()
                                img.save(img_buf, format="JPEG", quality=95)
                                pdf_page = new_doc.new_page(width=img.width, height=img.height)
                                pdf_page.insert_image(pdf_page.rect, stream=img_buf.getvalue())
                        merged_pdf_bytes_io = io.BytesIO()
                        new_doc.save(merged_pdf_bytes_io)
                        new_doc.close()
                        merged_pdf_bytes = merged_pdf_bytes_io.getvalue()
                    
                    is_jpeg = final_name.lower().endswith(('.jpg', '.jpeg'))
                    orig_total_bytes = sum([len(p["file_bytes"]) for p in group_pages])
                    
                    if is_jpeg and len(indices) == 1:
                        p_data = group_pages[0]
                        if "pdf" in p_data["mime_type"].lower():
                            src_doc = fitz.open(stream=p_data["file_bytes"], filetype="pdf")
                            page = src_doc.load_page(p_data["pdf_page_idx"])
                            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
                            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                            src_doc.close()
                        else:
                            img = Image.open(io.BytesIO(p_data["file_bytes"]))
                            img = ImageOps.exif_transpose(img)
                            
                        try:
                            rot = int(rotations.get(str(p_data["global_idx"]), 0))
                            if rot != 0: img = img.rotate(-rot, expand=True)
                        except: pass
                        
                        img_buf = io.BytesIO()
                        img.save(img_buf, format="JPEG", quality=95)
                        comp_bytes, out_mime = process_and_compress_file(img_buf.getvalue(), "image/jpeg", final_name)
                    else:
                        comp_bytes, out_mime = process_and_compress_file(merged_pdf_bytes, "application/pdf", final_name)
                        
                    orig_kb = orig_total_bytes / 1024
                    comp_kb = len(comp_bytes) / 1024
                    
                    src_display = ", ".join(source_names)
                    if len(src_display) > 30: src_display = src_display[:27] + "..."
                    
                    results.append({
                        "original_name": f"분석결과 ({src_display})",
                        "suggested_filename": final_name,
                        "category": doc_info.get("doc_category", "기타"),
                        "client_name": doc_info.get("client_name", ""),
                        "mime": out_mime,
                        "orig_kb": orig_kb,
                        "comp_kb": comp_kb,
                        "bytes": comp_bytes,
                        "is_unclassified": is_unclassified
                    })
                
                except Exception as e:
                    failed_docs.append(final_name)
                
                progress_bar.progress(min((idx + 1) * progress_step, 1.0))
                
            if failed_docs:
                status_text.warning("일부 서류 처리에 실패했지만, 나머지 서류의 최적화는 완료되었습니다.")
                st.warning("⚠️ **아래 서류는 파일 손상 또는 변환 중 오류가 발생하여 제외되었습니다. 원본 파일을 직접 확인해 주세요:**\n\n" + "\n".join([f"- {f}" for f in failed_docs]))
            else:
                status_text.success("모든 서류의 묶기/분할 및 스마트 회전 최적화가 완료되었습니다.")
                
            st.session_state.analysis_results = results

    if st.session_state.analysis_results:
        st.markdown("---")
        st.subheader("변환 완료된 서류 다운로드")
        
        zip_buffer = io.BytesIO()
        final_downloads = []
        
        for idx, item in enumerate(st.session_state.analysis_results):
            col1, col2, col3 = st.columns([3, 3, 2])
            
            with col1:
                st.write(f"**출처**: `{item['original_name']}`")
                st.caption(f"{item['orig_kb']:.1f} KB ➡️ **{item['comp_kb']:.1f} KB**")
                if item.get("is_unclassified"):
                    st.warning("⚠️ 규칙 미확인 서류 (파일명 수동 확인 필요)")
                
            with col2:
                user_edited_name = st.text_input(
                    "파일명", 
                    value=item['suggested_filename'], 
                    key=f"edit_{idx}",
                    label_visibility="collapsed"
                )
                final_downloads.append((user_edited_name, item['bytes'], item['mime']))
                
            with col3:
                st.download_button("⬇️ 개별 다운로드", data=item['bytes'], file_name=user_edited_name, mime=item['mime'], key=f"dl_btn_{idx}")
                
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for fname, fbytes, _ in final_downloads:
                zip_file.writestr(fname, fbytes)
                
        zip_buffer.seek(0)
        today_str = datetime.date.today().strftime("%Y%m%d")
        
        st.markdown("---")
        st.download_button(
            "📦 전체 서류 ZIP 다운로드",
            data=zip_buffer,
            file_name=f"CRM_Documents_{today_str}.zip",
            mime="application/zip",
            type="primary",
            use_container_width=True
        )

        st.markdown("---")
        if st.button("🔄 전체 리셋", type="secondary", use_container_width=True):
            for key in list(st.session_state.keys()):
                if key != "password_correct":
                    del st.session_state[key]
            st.session_state.uploader_key = str(uuid.uuid4())
            st.rerun()

# ------------------------------------------
# 메뉴 5: 잡오퍼 DOCX 생성기 (DEV)
# ------------------------------------------
elif app_mode == MENU_5:
    st.title(MENU_5)
    st.caption("손님 여권, 채용 공고(또는 기존 잡오퍼) 및 고용주 정보를 조합하여 캐나다 주별 오버타임 조항이 적용된 MS Word(.docx) 잡오퍼를 생성합니다.")
    
    if "job_offer_data" not in st.session_state:
        st.session_state.job_offer_data = {}
        
    doc_mode = st.radio("작성 모드를 선택하세요", ["🆕 신규 잡오퍼 생성", "🔄 기존 잡오퍼 연장/업데이트"])
    
    st.markdown("---")
    
    # 1. 손님 여권 업로드
    st.subheader("1. 손님 정보 (여권 업로드)")
    passport_file = st.file_uploader("손님 여권 이미지 또는 PDF", type=['jpg', 'jpeg', 'png', 'pdf'], key="jo_passport")
    
    # 2. 연장 모드일 경우 기존 잡오퍼 업로드
    existing_parsed = {}
    if doc_mode == "🔄 기존 잡오퍼 연장/업데이트":
        st.subheader("2. 기존 잡오퍼 서류 (업로드 시 기존 정보 자동 로드)")
        old_jo_file = st.file_uploader("기존 잡오퍼 (DOCX 또는 PDF)", type=['pdf', 'docx'], key="jo_old_file")
        if old_jo_file and st.button("기존 잡오퍼 분석하여 정보 가져오기"):
            with st.spinner("기존 잡오퍼 분석 중..."):
                file_b = old_jo_file.getvalue()
                m_type = old_jo_file.type if old_jo_file.type else "application/pdf"
                existing_parsed = parse_existing_job_offer(file_b, m_type)
                if existing_parsed:
                    st.success("기존 잡오퍼 정보를 성공적으로 읽어왔습니다! 아래 필드에 자동 적용됩니다.")
                    st.session_state.job_offer_data.update(existing_parsed)

    # 3. 채용 공고 분석 (URL 크롤링 연동)
    st.subheader("3. 채용 공고 (광고 링크 또는 내용 붙여넣기)")
    job_posting_text = st.text_area("채용 공고 링크(URL) 또는 공고 텍스트를 입력하세요", height=100, placeholder="https://www.jobspider.com/job/... 또는 공고 본문 텍스트")
    logo_file = st.file_uploader("회사 로고 이미지 (선택 사항)", type=['jpg', 'jpeg', 'png'], key="jo_logo")
    
    if st.button("AI 채용공고 & 여권 분석 시작"):
        with st.spinner("채용 공고 웹페이지 실시간 접속 및 서류 정보 분석 중..."):
            extracted_info = {}
            
            # 로고 이미지 바이트 저장
            if logo_file:
                extracted_info['logo_bytes'] = logo_file.getvalue()

            # 여권 분석
            if passport_file:
                pass_img = process_uploaded_file_to_image(passport_file)
                pass_data = extract_imm5476_info(pass_img)
                if pass_data:
                    extracted_info['client_name'] = format_full_name(pass_data.get('surname', ''), pass_data.get('given_name', ''))
                    extracted_info['client_dob'] = pass_data.get('dob', '')
            
            # 채용공고 분석 (URL 여부 판별 후 실시간 크롤링)
            if job_posting_text.strip():
                raw_input = job_posting_text.strip()
                is_url = bool(re.search(r'https?://[^\s]+|www\.[^\s]+', raw_input))
                
                if is_url:
                    url_match = re.search(r'(https?://[^\s]+|www\.[^\s]+)', raw_input).group(0)
                    fetched_text = fetch_url_content(url_match)
                    if fetched_text:
                        text_to_analyze = fetched_text
                    else:
                        text_to_analyze = raw_input
                else:
                    text_to_analyze = raw_input
                    
                prompt_job = f"""
                Analyze this job posting webpage/text content carefully:
                ---
                {text_to_analyze}
                ---

                Extract the following details into exact JSON:
                - employer_name: Full employer or company name including legal name and 'dba' if present (e.g. 'Agape Sushi Inc. dba Hiro Japan Sushi Xpress')
                - job_title: Position or Job Title (e.g. 'Food Service Supervisor')
                - wage: Hourly wage rate in numerical string format (e.g. '20.15')
                - hours: Working hours per week (e.g. '30-40')
                - job_location: Exact work location address including suite, street, city, province, postal code (e.g. '1644 Hillside Ave Suite 46A, Victoria, BC V8T 2C5')
                - employer_address: Corporate or Employer address if different, otherwise same as job_location
                - employer_phone: Contact telephone number if present
                - employer_email: Contact email address if present
                - benefits: Benefits (e.g. '4% vacation pay')
                - job_duties: Array of bullet point job duty strings

                Return ONLY raw valid JSON object without markdown formatting.
                """
                try:
                    resp = safe_generate_content([prompt_job])
                    clean = resp.text.strip().replace('```json', '').replace('```', '')
                    job_extracted = json.loads(clean)
                    extracted_info.update(job_extracted)
                except Exception as e:
                    st.warning(f"채용공고 분석 중 일부 오류: {e}")
                    
            st.session_state.job_offer_data.update(extracted_info)
            st.success("채용공고 웹페이지 정보가 정확하게 추출되어 아래 입력창에 적용되었습니다.")

    st.markdown("---")
    st.subheader("4. 최종 잡오퍼 정보 확인 및 수정")
    
    jo_data = st.session_state.job_offer_data
    
    col_c1, col_c2 = st.columns(2)
    with col_c1:
        c_name = st.text_input("손님 영문 성명 (Client Name)", value=jo_data.get('client_name', ''))
        offer_dt = st.date_input("오퍼 작성일 (Offer Date)", datetime.date.today()).strftime("%B %d, %Y")
        term_str = st.text_input("계약 기간 (Term)", value=jo_data.get('employment_term', '3-year'))
    with col_c2:
        c_dob = st.text_input("손님 생년월일 (Client DOB)", value=jo_data.get('client_dob', ''))
        start_dt_str = st.text_input("근무 시작일 (Start Date)", value=jo_data.get('start_date', 'The employment start date will be as soon as possible upon the employee’s authorization to work in Canada.'))

    st.markdown("#### 🏢 고용주 및 회사 정보 (미입력 시 기존 정보 유지)")
    col_e1, col_e2 = st.columns(2)
    with col_e1:
        emp_name = st.text_input("회사명 (Employer / Company Name)", value=jo_data.get('employer_name', ''))
        signer_n = st.text_input("대표자/서명자 성명 (Signer Name)", value=jo_data.get('signer_name', ''))
        signer_t = st.text_input("대표자 직책 (Signer Title)", value=jo_data.get('signer_title', 'Director'))
    with col_e2:
        emp_addr = st.text_input("회사 대표 주소 (Employer Address)", value=jo_data.get('employer_address', ''))
        emp_phone = st.text_input("회사 전화번호 (Employer Phone)", value=jo_data.get('employer_phone', ''))
        emp_email = st.text_input("회사 이메일 (Employer Email)", value=jo_data.get('employer_email', ''))

    st.markdown("#### 💼 근무 조건 및 직무")
    col_j1, col_j2 = st.columns(2)
    with col_j1:
        j_title = st.text_input("직책 (Job Title)", value=jo_data.get('job_title', ''))
        j_wage = st.text_input("시급 (Hourly Wage, CAD)", value=str(jo_data.get('wage', '20.15')))
        j_hours = st.text_input("주당 근무시간 (Weekly Hours)", value=str(jo_data.get('hours', '30-40')))
    with col_j2:
        j_loc = st.text_input("실제 근무지 주소 (Job Location)", value=jo_data.get('job_location', emp_addr))
        j_benefits = st.text_input("혜택 (Benefits)", value=jo_data.get('benefits', '4% vacation pay'))

    # 오버타임 조항 주(Province)별 자동 추천
    auto_ot_clause = get_provincial_overtime_clause(j_loc if j_loc else emp_addr)
    j_ot = st.text_area("오버타임 조항 (근무지 주소에 따라 자동 계산됨)", value=auto_ot_clause, height=80)

    duties_input_str = jo_data.get('job_duties', [])
    if isinstance(duties_input_str, list):
        duties_input_str = "\n".join(duties_input_str)
        
    j_duties_text = st.text_area("주요 직무 (Job Duties - 한 줄에 하나씩 입력)", value=duties_input_str, height=150)

    st.markdown("---")
    if st.button("📄 MS Word (.docx) 잡오퍼 생성 및 다운로드", type="primary", use_container_width=True):
        if not c_name or not emp_name or not j_title:
            st.error("손님 성명, 회사명, 직책은 필수 입력 항목입니다.")
        else:
            final_jo_dict = {
                "client_name": c_name,
                "client_dob": c_dob,
                "offer_date": offer_dt,
                "employment_term": term_str,
                "start_date": start_dt_str,
                "employer_name": emp_name,
                "signer_name": signer_n,
                "signer_title": signer_t,
                "employer_address": emp_addr,
                "employer_phone": emp_phone,
                "employer_email": emp_email,
                "job_title": j_title,
                "wage": j_wage,
                "hours": j_hours,
                "job_location": j_loc,
                "benefits": j_benefits,
                "overtime_clause": j_ot,
                "job_duties": j_duties_text,
                "logo_bytes": jo_data.get('logo_bytes')
            }
            
            docx_bytes = generate_job_offer_docx(final_jo_dict)
            
            crm_client = "NAME"
            if c_name:
                if re.search(r'[가-힣]', c_name):
                    crm_client = c_name.replace(" ", "")
                else:
                    parts = c_name.strip().split()
                    if parts: crm_client = parts[0].capitalize()
                    
            out_filename = f"[Job Offer]_{crm_client}.docx"
            
            st.success("잡오퍼 DOCX 문서가 성공적으로 생성되었습니다!")
            st.download_button(
                label="📥 Job Offer .docx 파일 다운로드",
                data=docx_bytes,
                file_name=out_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
                use_container_width=True
            )
            
    st.markdown("---")
    if st.button("🔄 전체 리셋", type="secondary", use_container_width=True, key="reset_btn_m5"):
        st.session_state.job_offer_data = {}
        st.rerun()
